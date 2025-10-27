"""Cover Letter Generator Module"""

import os
import re
import time
import json
from pathlib import Path
from typing import Optional, List, Dict
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
import pythoncom
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select
from .utils import (
    TIMEOUT, PAGE_LOAD,
    navigate_to_folder, get_pagination_pages, go_to_next_page,
    get_jobs_from_page, sanitize_filename
)
from .config import load_app_config
from .agents import AgentFactory


class CoverLetterGenerator:
    """Generate and manage cover letters for job applications"""

    def __init__(self, driver, resume_text, prompt_template, api_key=None, cover_letters_folder: Optional[str] = None):
        self.driver = driver
        self.resume_text = resume_text
        self.prompt_template = prompt_template  # Keep for backwards compatibility
        
        # Load config
        config = load_app_config()
        
        if cover_letters_folder is None:
            cover_letters_folder = config.get("paths", {}).get("cover_letters_dir", "cover_letters")
        
        self.cover_letters_dir = Path(cover_letters_folder)
        self.cover_letters_dir.mkdir(exist_ok=True)
        
        # Initialize agent factory
        agent_config = {
            "cover_letter_agent": {
                "provider": config.agents.cover_letter_agent.get("provider", "gemini"),
                "model": config.agents.cover_letter_agent.get("model", "gemini-1.5-flash")
            }
        }
        
        self.factory = AgentFactory(
            config=agent_config,
            enable_tracking=config.agents.enable_token_tracking
        )
        
        self.agent = self.factory.get_cover_letter_agent()
        print(f"✅ CoverLetterGenerator initialized with {self.agent.provider}/{self.agent.model}")
        
    def get_document_name(self, company, job_title):
        company_clean = sanitize_filename(company)
        title_clean = sanitize_filename(job_title)
        return f"{company_clean}_{title_clean}"
    
    def cover_letter_exists(self, company, job_title):
        doc_name = self.get_document_name(company, job_title)
        pdf_path = self.cover_letters_dir / f"{doc_name}.pdf"
        return pdf_path.exists()
    
    def parse_jobs_from_page(self):
        jobs = []
        try:
            job_rows = get_jobs_from_page(self.driver)
            print(f"   Found {len(job_rows)} job listings on this page")
            
            for idx, row in enumerate(job_rows, 1):
                try:
                    cells = row.find_elements(By.TAG_NAME, "td")
                    if len(cells) < 4:
                        continue
                    
                    job_title_elem = cells[0].find_element(By.TAG_NAME, "a")
                    job_title = job_title_elem.text.strip()
                    job_id = job_title_elem.get_attribute("href").split("=")[-1]
                    company = cells[1].text.strip()
                    
                    jobs.append({
                        "job_id": job_id,
                        "job_title": job_title,
                        "company": company,
                        "title_element": job_title_elem,
                        "row_index": idx
                    })
                except Exception as e:
                    print(f"   ⚠ Error extracting job {idx}: {e}")
                    continue
            
            return jobs
        except Exception as e:
            print(f"   ✗ Error getting jobs from page: {e}")
            return []
    
    def get_job_details_from_cache(self, job_id: str, cached_jobs: List[Dict]) -> Optional[Dict]:
        if not cached_jobs:
            return None
        
        # Simple ID matching
        for job in cached_jobs:
            if job.get("id") == job_id or job.get("job_id") == job_id:
                print(f"      ✓ Found in cache")
                
                # Build description from key fields
                desc_parts = []
                if job.get("summary") and job["summary"] != "N/A":
                    desc_parts.append(f"Job Summary:\n{job['summary']}")
                if job.get("responsibilities") and job["responsibilities"] != "N/A":
                    desc_parts.append(f"\nResponsibilities:\n{job['responsibilities']}")
                if job.get("skills") and job["skills"] != "N/A":
                    desc_parts.append(f"\nRequired Skills:\n{job['skills']}")
                if job.get("additional_info") and job["additional_info"] != "N/A":
                    desc_parts.append(f"\nAdditional Info:\n{job['additional_info']}")
                
                description = "\n".join(desc_parts)
                
                return {
                    "job_id": job_id,
                    "job_title": job.get("title", ""),
                    "company": job.get("company", ""),
                    "description": description
                }
        
        return None  # Not in cache
    
    def scrape_and_cache_job(self, job_basic_info: Dict, cached_jobs_path: str) -> Optional[Dict]:
        from .scraper import WaterlooWorksScraper
        
        print(f"      🔍 Not in cache - scraping live...")
        
        try:
            # Create a temporary scraper to get details
            scraper = WaterlooWorksScraper(self.driver)
            
            # Add row_element for scraper to click
            job_basic_info["row_element"] = job_basic_info.get("title_element")
            
            # Scrape full details
            job_data = scraper.get_job_details(job_basic_info)
            
            if not job_data:
                print(f"      ✗ Failed to scrape job details")
                return None
            
            # Add to cache file
            cached_jobs = []
            if os.path.exists(cached_jobs_path):
                with open(cached_jobs_path, 'r', encoding='utf-8') as f:
                    cached_jobs = json.load(f)
            
            cached_jobs.append(job_data)
            
            # Save updated cache
            with open(cached_jobs_path, 'w', encoding='utf-8') as f:
                json.dump(cached_jobs, f, indent=2, ensure_ascii=False)
            
            print(f"      ✓ Scraped and cached job details")
            
            # Build description for cover letter generation
            desc_parts = []
            if job_data.get("summary") and job_data["summary"] != "N/A":
                desc_parts.append(f"Job Summary:\n{job_data['summary']}")
            if job_data.get("responsibilities") and job_data["responsibilities"] != "N/A":
                desc_parts.append(f"\nResponsibilities:\n{job_data['responsibilities']}")
            if job_data.get("skills") and job_data["skills"] != "N/A":
                desc_parts.append(f"\nRequired Skills:\n{job_data['skills']}")
            if job_data.get("additional_info") and job_data["additional_info"] != "N/A":
                desc_parts.append(f"\nAdditional Info:\n{job_data['additional_info']}")
            
            description = "\n".join(desc_parts)
            
            return {
                "job_id": job_data.get("id"),
                "job_title": job_data.get("title", ""),
                "company": job_data.get("company", ""),
                "description": description
            }
            
        except Exception as e:
            print(f"      ✗ Error scraping job: {e}")
            return None
    
    def generate_cover_letter_text(self, company, job_title, description):
        try:
            # Use the agent's generate method
            result = self.agent.generate_cover_letter(
                company=company,
                job_title=job_title,
                job_description=description,
                resume_text=self.resume_text,
                max_retries=3
            )
            
            if result:
                print(f"      ✓ Generated {len(result.split())} word cover letter")
                return result
            else:
                print(f"      ✗ Failed to generate cover letter")
                return None
                
        except Exception as e:
            print(f"      ✗ Error generating cover letter: {str(e)[:100]}")
            return None
    
    def save_cover_letter(self, company, job_title, cover_text, template_path="template.docx"):
        doc_name = self.get_document_name(company, job_title)
        
        try:
            # Load template
            document = Document(template_path)
            
            # Add generated text
            generated_text = document.add_paragraph().add_run(
                f"""Dear Hiring Manager,
        
{cover_text}

Aman Zaveri"""
            )
            
            generated_text.font.size = Pt(11)
            generated_text.font.name = "Garamond"
            
            # Save as DOCX first
            docx_path = self.cover_letters_dir / f"{doc_name}.docx"
            document.save(str(docx_path))
            
            # Convert to PDF
            pythoncom.CoInitialize()
            convert(str(docx_path))
            pythoncom.CoUninitialize()
            
            # Remove DOCX file
            docx_path.unlink()
            
            print(f"      ✓ Saved: {doc_name}.pdf")
            return True
            
        except Exception as e:
            print(f"      ✗ Error saving cover letter: {e}")
            return False
    
    def generate_all_cover_letters(self, folder_name: str, cached_jobs_path=None):
        if not folder_name:
            raise ValueError(
                "folder_name is required for cover letter generation. "
                "You must specify which WaterlooWorks folder to generate cover letters for. "
                "This prevents accidentally generating cover letters for all 2000+ jobs."
            )
        
        stats = {
            "total_jobs": 0,
            "skipped_existing": 0,
            "generated": 0,
            "failed": 0
        }
        
        # Load cached jobs if available
        cached_jobs = None
        if cached_jobs_path and os.path.exists(cached_jobs_path):
            try:
                with open(cached_jobs_path, 'r', encoding='utf-8') as f:
                    cached_jobs = json.load(f)
                print(f"✓ Loaded {len(cached_jobs)} cached jobs")
            except Exception as e:
                print(f"⚠ Could not load cached jobs: {e}")
        
        # Navigate to specified folder
        print(f"\n📁 Navigating to '{folder_name}' folder...")
        if not navigate_to_folder(self.driver, folder_name):
            return stats
        print(f"   ✓ Successfully navigated to '{folder_name}' folder")
        
        # Get number of pages
        num_pages = get_pagination_pages(self.driver)
        print(f"\n📄 Found {num_pages} page(s) of jobs")
        
        # Collect jobs from all pages
        all_jobs = []
        for page in range(1, num_pages + 1):
            print(f"\n📊 Extracting jobs from page {page}/{num_pages}...")
            
            # Get jobs from current page
            jobs = self.parse_jobs_from_page()
            all_jobs.extend(jobs)
            
            print(f"   ✓ Extracted {len(jobs)} jobs from page {page}")
            
            # Go to next page if not the last one
            if page < num_pages:
                print(f"   ➡️  Going to page {page + 1}...")
                go_to_next_page(self.driver)
        
        stats["total_jobs"] = len(all_jobs)
        
        if not all_jobs:
            print(f"No jobs found in '{folder_name}' folder")
            return stats
        
        print(f"\n🎯 Processing {len(all_jobs)} jobs total...")
        
        # Process each job
        for idx, job_basic in enumerate(all_jobs, 1):
            company = job_basic["company"]
            job_title = job_basic["job_title"]
            job_id = job_basic["job_id"]
            
            print(f"\n[{idx}/{len(all_jobs)}] {company} - {job_title}")
            
            # Check if cover letter already exists
            if self.cover_letter_exists(company, job_title):
                print(f"      ⏭ Cover letter already exists, skipping")
                stats["skipped_existing"] += 1
                continue
            
            # Try to get from cache first
            job_details = self.get_job_details_from_cache(job_id, cached_jobs or [])
            
            # If not in cache, scrape it and add to cache
            if not job_details and cached_jobs_path:
                job_details = self.scrape_and_cache_job(job_basic, cached_jobs_path)
            
            if not job_details:
                print(f"      ✗ Could not get job details")
                stats["failed"] += 1
                continue
            
            # Check if we have a real description
            description = job_details.get("description", "")
            if not description or len(description) < 50:
                print(f"      ⏭ Skipping (description too short or missing)")
                stats["failed"] += 1
                continue
            
            # Generate cover letter text
            print(f"      🤖 Generating cover letter...")
            cover_text = self.generate_cover_letter_text(
                company,
                job_title,
                description
            )
            
            if not cover_text:
                stats["failed"] += 1
                continue
            
            # Save cover letter
            if self.save_cover_letter(company, job_title, cover_text):
                stats["generated"] += 1
            else:
                stats["failed"] += 1
        
        return stats


class CoverLetterUploader:
    
    def __init__(self, driver, cover_letters_folder: Optional[str] = None):
        self.driver = driver
        
        if cover_letters_folder is None:
            from .config import load_app_config
            config = load_app_config()
            cover_letters_folder = config.get("paths", {}).get("cover_letters_dir", "cover_letters")
        
        self.cover_letters_dir = Path(cover_letters_folder)
    
    def navigate_to_upload_menu(self):
        try:
            print("\n📤 Navigating to upload menu...")
            
            # Go to home page
            home_button = WebDriverWait(self.driver, TIMEOUT).until(
                EC.element_to_be_clickable((By.ID, "outerTemplateTabs_overview"))
            )
            home_button.click()
            time.sleep(2)
            
            # Open main menu
            # menu_button = WebDriverWait(self.driver, TIMEOUT).until(
            #     EC.element_to_be_clickable(
            #         (By.CSS_SELECTOR, "[aria-label='Toggle Main Menu']")
            #     )
            # )
            # menu_button.click()
            # time.sleep(1)
            
            # Click upload documents
            upload_button = WebDriverWait(self.driver, TIMEOUT).until(
                EC.element_to_be_clickable(
                    (By.CSS_SELECTOR, "[data-pt-classes='tip--default']")
                )
            )
            upload_button.click()
            time.sleep(1)
            
            # Click internal upload button
            upload_buttons = WebDriverWait(self.driver, TIMEOUT).until(
                EC.presence_of_all_elements_located(
                    (By.CSS_SELECTOR, "[class='btn__default--text btn--info  display--flex align--middle']")
                )
            )
            upload_buttons[0].click()
            time.sleep(1)
            
            print("   ✓ Navigated to upload menu")
            return True
            
        except Exception as e:
            print(f"   ✗ Error navigating to upload menu: {e}")
            return False
    
    def upload_file(self, pdf_filename):
        try:
            # Extract document name (without .pdf)
            doc_name = pdf_filename.replace(".pdf", "")
            
            # Write cover letter name
            name_input = WebDriverWait(self.driver, TIMEOUT).until(
                EC.presence_of_element_located((By.ID, "docName"))
            )
            name_input.clear()
            name_input.send_keys(doc_name)
            
            # Select document type (66 = Cover Letter)
            select_element = self.driver.find_element(By.ID, "docType")
            select = Select(select_element)
            select.select_by_value("66")
            
            # Upload file
            file_path = str((self.cover_letters_dir / pdf_filename).absolute())
            file_input = WebDriverWait(self.driver, TIMEOUT).until(
                EC.presence_of_element_located((By.ID, "fileUpload_docUpload"))
            )
            file_input.send_keys(file_path)
            time.sleep(3)  # Wait for file to upload
            
            # Submit
            submit_button = WebDriverWait(self.driver, TIMEOUT).until(
                EC.element_to_be_clickable((By.ID, "submitFileUploadFormBtn"))
            )
            submit_button.click()
            time.sleep(2)
            
            print(f"      ✓ Uploaded: {doc_name}")
            return True
            
        except Exception as e:
            print(f"      ✗ Error uploading {pdf_filename}: {e}")
            return False
    
    def get_uploaded_files_log(self):
        from .config import load_app_config
        config = load_app_config()
        data_dir = config.get("paths", {}).get("data_dir", "data")
        return Path(data_dir) / "uploaded_cover_letters.json"
    
    def load_uploaded_files(self):
        log_file = self.get_uploaded_files_log()
        if log_file.exists():
            with open(log_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return set(data.get("uploaded_files", []))
        return set()
    
    def save_uploaded_file(self, filename):
        log_file = self.get_uploaded_files_log()
        log_file.parent.mkdir(exist_ok=True)
        
        # Load existing
        uploaded = self.load_uploaded_files()
        uploaded.add(filename)
        
        # Save updated list
        with open(log_file, 'w', encoding='utf-8') as f:
            json.dump({"uploaded_files": sorted(list(uploaded))}, f, indent=2)
    
    def upload_all_cover_letters(self):
        stats = {
            "total_files": 0,
            "uploaded": 0,
            "skipped_existing": 0,
            "failed": 0
        }
        
        # Get all PDF files
        pdf_files = list(self.cover_letters_dir.glob("*.pdf"))
        stats["total_files"] = len(pdf_files)
        
        if not pdf_files:
            print("No cover letters found to upload")
            return stats
        
        # Load already uploaded files
        uploaded_files = self.load_uploaded_files()
        
        # Filter out already uploaded files
        files_to_upload = [f for f in pdf_files if f.name not in uploaded_files]
        stats["skipped_existing"] = len(pdf_files) - len(files_to_upload)
        
        print(f"\n📤 Found {len(pdf_files)} total cover letters")
        if stats["skipped_existing"] > 0:
            print(f"   ⏭️  Skipping {stats['skipped_existing']} already uploaded")
        print(f"   📤 Uploading {len(files_to_upload)} new files")
        
        if not files_to_upload:
            print("\n✅ All cover letters already uploaded!")
            return stats
        
        # Navigate to upload menu once
        if not self.navigate_to_upload_menu():
            return stats
        
        # Upload each file
        for idx, pdf_path in enumerate(files_to_upload, 1):
            print(f"\n[{idx}/{len(files_to_upload)}] Uploading {pdf_path.name}...")
            
            if self.upload_file(pdf_path.name):
                stats["uploaded"] += 1
                # Mark as uploaded
                self.save_uploaded_file(pdf_path.name)
                
                # Navigate back to upload page for next file (except last one)
                if idx < len(files_to_upload):
                    if not self.navigate_to_upload_menu():
                        break
            else:
                stats["failed"] += 1
        
        return stats
