"""
Waterloo Works Automator - Main Pipeline
Scrapes jobs, analyzes matches, and shows best opportunities
"""

import json
import os
import re
import getpass
from datetime import datetime
from typing import List, Dict
from pathlib import Path
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
import pythoncom

# Load environment variables
load_dotenv()

from modules.auth import WaterlooWorksAuth
from modules.scraper import WaterlooWorksScraper
from modules.matcher import ResumeMatcher, load_config


class JobAnalyzer:
    """Main pipeline for scraping and analyzing WaterlooWorks jobs"""
    
    def __init__(self, config_path: str = "config.json", input_folder: str = "input", 
                 cover_letters_folder: str = "cover_letters", waterlooworks_folder: str = None):
        """
        Initialize analyzer with configuration
        
        Args:
            config_path: Path to config.json
            input_folder: Folder containing resume
            cover_letters_folder: Folder to save/read cover letters
            waterlooworks_folder: WaterlooWorks folder name (overrides config)
        """
        self.config = load_config(config_path)
        self.matcher = ResumeMatcher(config_path)
        self.auth = None  # Will be set during scraping
        self.scraper = None  # Will be set during scraping
        
        # Custom folder paths
        self.input_folder = input_folder
        self.cover_letters_folder = cover_letters_folder
        self.waterlooworks_folder = waterlooworks_folder or self.config.get("waterlooworks_folder", "geese")
        
        print("✅ Job Analyzer initialized")
        print(f"   📂 Input folder: {self.input_folder}/")
        print(f"   📄 Cover letters folder: {self.cover_letters_folder}/")
        print(f"   🗂️  WaterlooWorks folder: {self.waterlooworks_folder}")
        print()
    
    def run_realtime_pipeline(self, auto_save_to_folder: bool = True) -> List[Dict]:
        """
        Run REAL-TIME pipeline: scrape → analyze → save IMMEDIATELY per job
        
        Shows score for each job as it's processed and saves high-scoring jobs instantly.
        
        Args:
            auto_save_to_folder: If True, save high-scoring jobs to WW folder immediately
        
        Returns:
            List of analyzed jobs sorted by fit score
        """
        print("=" * 70)
        print("🚀 WATERLOO WORKS REAL-TIME JOB ANALYZER")
        print("=" * 70)
        print()
        
        # Get configuration
        auto_save_threshold = self.config.get("matcher", {}).get("auto_save_threshold", 30)
        folder_name = self.waterlooworks_folder  # Use custom folder
        preferred_locations = self.config.get("preferred_locations", [])
        keywords = self.config.get("keywords_to_match", [])
        avoid_companies = self.config.get("companies_to_avoid", [])
        
        print(f"⚙️  Auto-save threshold: {auto_save_threshold}")
        print(f"📁 Target folder: '{folder_name}'")
        print()
        
        # Step 1: Login and navigate
        print("🔐 Logging into WaterlooWorks...")
        username = input("Username (UW email): ").strip()
        import getpass
        password = getpass.getpass("Password: ")
        print()
        
        from modules.auth import WaterlooWorksAuth
        auth = WaterlooWorksAuth(username, password)
        auth.login()
        self.auth = auth
        
        llm_provider = self.config.get("matcher", {}).get("llm_provider", "gemini")
        scraper = WaterlooWorksScraper(auth.driver, llm_provider=llm_provider)
        self.scraper = scraper
        scraper.go_to_jobs_page()
        
        # Step 2: Process jobs in real-time
        all_results = []
        saved_count = 0
        skipped_count = 0
        total_jobs = 0
        
        try:
            # Get pagination info
            try:
                num_pages = scraper.get_pagination_pages()
                print(f"📄 Total pages to process: {num_pages}\n")
            except Exception:
                num_pages = 1
                print("📄 Single page detected\n")
            
            print("=" * 70)
            print("🔄 PROCESSING JOBS IN REAL-TIME")
            print("=" * 70)
            print()
            
            # Process each page
            for page in range(1, num_pages + 1):
                print(f"📄 Page {page}/{num_pages}")
                print("-" * 70)
                
                rows = scraper.get_job_table()
                
                for i, row in enumerate(rows, 1):
                    total_jobs += 1
                    
                    # Parse basic job info
                    job_data = scraper.parse_job_row(row)
                    if not job_data or not job_data.get('id'):
                        continue
                    
                    # Show job being processed
                    print(f"\n[Job {total_jobs}] {job_data.get('title', 'Unknown')} @ {job_data.get('company', 'N/A')}")
                    print(f"           📍 {job_data.get('location', 'N/A')}")
                    
                    # Get detailed job info
                    print(f"           🔍 Scraping details...", end=" ")
                    job_data = scraper.get_job_details(job_data)
                    print("✓")
                    
                    # Analyze match IMMEDIATELY
                    print(f"           🧠 Analyzing match...", end=" ")
                    result = self.matcher.analyze_single_job(job_data)
                    print("✓")
                    
                    # Extract score
                    match = result["match"]
                    fit_score = match["fit_score"]
                    
                    # Show score with color coding
                    if fit_score >= 70:
                        emoji = "🟢"
                    elif fit_score >= 50:
                        emoji = "🟡"
                    elif fit_score >= 30:
                        emoji = "🟠"
                    else:
                        emoji = "🔴"
                    
                    print(f"           {emoji} FIT SCORE: {fit_score}/100")
                    
                    # Apply filters
                    should_skip = False
                    
                    # Location filter
                    if preferred_locations:
                        job_location = job_data.get("location", "").lower()
                        if not any(loc.lower() in job_location for loc in preferred_locations):
                            print(f"           ⏭️  Skipped (location filter)")
                            should_skip = True
                    
                    # Keyword filter
                    if not should_skip and keywords:
                        job_text_parts = []
                        for field in ['title', 'summary', 'responsibilities', 'skills']:
                            if job_data.get(field) and job_data[field] != 'N/A':
                                job_text_parts.append(job_data[field])
                        job_text = " ".join(job_text_parts).lower()
                        
                        if not any(kw.lower() in job_text for kw in keywords):
                            print(f"           ⏭️  Skipped (keyword filter)")
                            should_skip = True
                    
                    # Company filter
                    if not should_skip and avoid_companies:
                        company = job_data.get("company", "").lower()
                        if any(avoid.lower() in company for avoid in avoid_companies):
                            print(f"           ⏭️  Skipped (company filter)")
                            should_skip = True
                    
                    # Save if meets threshold
                    if not should_skip and fit_score >= auto_save_threshold and auto_save_to_folder:
                        print(f"           💾 Saving to '{folder_name}' folder...", end=" ")
                        
                        # Re-add row_element for saving
                        job_data["row_element"] = row
                        success = scraper.save_job_to_folder(job_data, folder_name=folder_name)
                        
                        if success:
                            print("✅ SAVED!")
                            saved_count += 1
                        else:
                            print("❌ Failed")
                            # Close panel if save failed
                            scraper.close_job_details_panel()
                        
                        # Remove row_element after saving
                        if "row_element" in job_data:
                            del job_data["row_element"]
                    elif not should_skip and fit_score < auto_save_threshold:
                        print(f"           ⏭️  Not saved (score < {auto_save_threshold})")
                        skipped_count += 1
                        # Close the panel since we're not saving
                        scraper.close_job_details_panel()
                    else:
                        skipped_count += 1
                        # Close the panel since we're skipping
                        scraper.close_job_details_panel()
                    
                    # Store result
                    all_results.append(result)
                
                # Go to next page
                if page < num_pages:
                    print(f"\n➡️  Moving to page {page + 1}...")
                    scraper.next_page()
                    print()
        
        except KeyboardInterrupt:
            print("\n\n⚠️  Process interrupted by user")
        except Exception as e:
            print(f"\n\n❌ Error during processing: {e}")
            import traceback
            traceback.print_exc()
        
        # Step 3: Save results to files
        print("\n" + "=" * 70)
        print("💾 SAVING RESULTS TO FILES")
        print("=" * 70)
        
        # Sort by fit score
        all_results.sort(key=lambda x: x["match"]["fit_score"], reverse=True)
        
        self._save_results(all_results)
        
        # Step 4: Show final summary
        print("\n" + "=" * 70)
        print("📊 FINAL SUMMARY")
        print("=" * 70)
        print(f"Total jobs processed: {total_jobs}")
        print(f"✅ Saved to WW folder: {saved_count}")
        print(f"⏭️  Skipped/Low score: {skipped_count}")
        print()
        
        self._show_summary(all_results[:10])  # Show top 10
        
        # Cleanup
        if self.auth:
            try:
                self.auth.close()
            except Exception:
                pass
        
        return all_results
    
    def run_full_pipeline(self, detailed: bool = True, force_rematch: bool = False, auto_save_to_folder: bool = False) -> List[Dict]:
        """
        Run complete pipeline: scrape → match → filter → save (+ optional auto-save to WW folder)
        
        Args:
            detailed: Whether to scrape detailed job info (slower but better)
            force_rematch: If True, ignore cached matches and recalculate all
            auto_save_to_folder: If True, automatically save high-scoring jobs to "geese" folder
        
        Returns:
            List of analyzed jobs sorted by fit score
        """
        print("=" * 70)
        print("🚀 WATERLOO WORKS JOB ANALYZER")
        print("=" * 70)
        print()
        
        # Step 1: Scrape jobs
        print("📥 Step 1: Scraping jobs from WaterlooWorks...")
        jobs = self._scrape_jobs(detailed=detailed)
        
        if not jobs:
            print("❌ No jobs found. Exiting.")
            return []
        
        print(f"✅ Found {len(jobs)} jobs\n")
        
        # Step 2: Analyze matches
        print("🔍 Step 2: Analyzing job matches...")
        if force_rematch:
            print("⚠️  Force rematch enabled - ignoring cache")
        results = self.matcher.batch_analyze(jobs, force_rematch=force_rematch)
        print(f"✅ Analyzed {len(results)} jobs\n")
        
        # Step 3: Filter results
        print("🎯 Step 3: Applying filters...")
        filtered_results = self._apply_filters(results)
        print(f"✅ {len(filtered_results)} jobs after filtering\n")
        
        # Step 4: Auto-save high-scoring jobs to WaterlooWorks folder (optional)
        if auto_save_to_folder and self.scraper:
            print("📁 Step 4: Auto-saving high-scoring jobs to WaterlooWorks folder...")
            self._auto_save_to_folder(filtered_results)
            print()
        
        # Step 5: Save results locally
        print("💾 Step 5: Saving results to local files...")
        self._save_results(filtered_results)
        print()
        
        # Step 6: Show summary
        self._show_summary(filtered_results)
        
        # Cleanup: Close browser if still open
        if self.auth:
            try:
                self.auth.close()
            except Exception:
                pass
        
        return filtered_results
    
    def _scrape_jobs(self, detailed: bool = True) -> List[Dict]:
        """Scrape jobs from WaterlooWorks with incremental saving"""
        try:
            # Load existing jobs from cache
            os.makedirs("data", exist_ok=True)
            cache_path = "data/jobs_scraped.json"
            existing_jobs = {}
            
            if os.path.exists(cache_path):
                print("📂 Loading existing jobs from cache...")
                with open(cache_path, 'r', encoding='utf-8') as f:
                    cached = json.load(f)
                    # Build a dict for fast lookup by job ID
                    existing_jobs = {job['id']: job for job in cached if 'id' in job}
                    print(f"   Found {len(existing_jobs)} cached jobs\n")
            
            # Get credentials
            print("🔐 Enter your WaterlooWorks credentials")
            username = input("Username (UW email): ").strip()
            password = getpass.getpass("Password: ")
            print()
            
            # Login
            auth = WaterlooWorksAuth(username, password)
            auth.login()
            
            # Store auth and scraper for later use (auto-save feature)
            self.auth = auth
            
            # Scrape with incremental saving
            llm_provider = self.config.get("matcher", {}).get("llm_provider", "gemini")
            scraper = WaterlooWorksScraper(auth.driver, llm_provider=llm_provider)
            self.scraper = scraper  # Store for later use
            scraper.go_to_jobs_page()
            
            # Pass existing jobs and save callback
            def save_callback(jobs):
                """Called after each page to save progress"""
                normalized = self._normalize_job_data(jobs)
                with open(cache_path, 'w', encoding='utf-8') as f:
                    json.dump(normalized, f, indent=2, ensure_ascii=False)
            
            jobs = scraper.scrape_all_jobs(
                include_details=detailed,
                existing_jobs=existing_jobs,
                save_callback=save_callback,
                save_every=5  # Save after every 5 new jobs (adjust as needed)
            )
            
            # DON'T close browser yet if we need to auto-save to folder
            # Browser will be closed at the end of the pipeline
            
            # Final normalization and save
            jobs = self._normalize_job_data(jobs)
            with open(cache_path, 'w', encoding='utf-8') as f:
                json.dump(jobs, f, indent=2, ensure_ascii=False)
            
            print(f"\n✅ Total jobs in cache: {len(jobs)}")
            print(f"   Newly scraped: {len([j for j in jobs if j.get('id') not in existing_jobs])}")
            print(f"   From cache: {len(existing_jobs)}\n")
            
            return jobs
        
        except Exception as e:
            print(f"❌ Error scraping jobs: {e}")
            
            # Try to use cached data
            if os.path.exists("data/jobs_scraped.json"):
                print("📂 Using cached jobs from previous run...")
                with open("data/jobs_scraped.json", 'r', encoding='utf-8') as f:
                    return json.load(f)
            
            return []
    
    def _normalize_job_data(self, jobs: List[Dict]) -> List[Dict]:
        """Normalize job data format for consistency"""
        # No normalization needed anymore - scraper already uses 'location'
        return jobs
    
    def _apply_filters(self, results: List[Dict]) -> List[Dict]:
        """Apply location, keyword, and company filters"""
        filtered = []
        
        preferred_locations = self.config.get("preferred_locations", [])
        keywords = self.config.get("keywords_to_match", [])
        avoid_companies = self.config.get("companies_to_avoid", [])
        min_score = self.config.get("matcher", {}).get("min_match_score", 0)
        
        for result in results:
            job = result["job"]
            match = result["match"]
            
            # Filter by minimum score
            if match["fit_score"] < min_score:
                continue
            
            # Filter by location
            if preferred_locations:
                job_location = job.get("location", "").lower()
                if not any(loc.lower() in job_location for loc in preferred_locations):
                    continue
            
            # Filter by keywords (if any match)
            if keywords:
                # Search across all text fields
                job_text_parts = []
                for field in ['title', 'summary', 'responsibilities', 'skills', 'employment_location_arrangement', 'work_term_duration']:
                    if job.get(field) and job[field] != 'N/A':
                        job_text_parts.append(job[field])
                job_text = " ".join(job_text_parts).lower()
                
                if not any(kw.lower() in job_text for kw in keywords):
                    continue
            
            # Filter out avoided companies
            if avoid_companies:
                company = job.get("company", "").lower()
                if any(avoid.lower() in company for avoid in avoid_companies):
                    continue
            
            filtered.append(result)
        
        return filtered
    
    def _auto_save_to_folder(self, results: List[Dict]):
        """
        Automatically save high-scoring jobs to WaterlooWorks folder
        
        Args:
            results: List of analyzed job results
        """
        if not self.scraper:
            print("  ⚠️  Scraper not available. Skipping auto-save.")
            return
        
        # Get thresholds from config
        auto_save_threshold = self.config.get("matcher", {}).get("auto_save_threshold", 50)
        folder_name = self.config.get("waterlooworks_folder", "geese")
        
        # Find jobs that meet the threshold
        jobs_to_save = [
            r for r in results 
            if r["match"]["fit_score"] >= auto_save_threshold
        ]
        
        if not jobs_to_save:
            print(f"  ℹ️  No jobs with fit score >= {auto_save_threshold}. Nothing to save.")
            return
        
        print(f"  🎯 Found {len(jobs_to_save)} jobs with fit score >= {auto_save_threshold}")
        print(f"  📁 Saving to WaterlooWorks folder: '{folder_name}'")
        print()
        
        saved_count = 0
        failed_count = 0
        
        for i, result in enumerate(jobs_to_save, 1):
            job = result["job"]
            score = result["match"]["fit_score"]
            
            print(f"  [{i}/{len(jobs_to_save)}] {job.get('title', 'Unknown')} - Score: {score}/100")
            
            success = self.scraper.save_job_to_folder(job, folder_name=folder_name)
            
            if success:
                saved_count += 1
            else:
                failed_count += 1
            
            print()
        
        print(f"  ✅ Successfully saved: {saved_count}/{len(jobs_to_save)}")
        if failed_count > 0:
            print(f"  ❌ Failed to save: {failed_count}/{len(jobs_to_save)}")
    
    def _save_results(self, results: List[Dict]):
        """Save analyzed results to JSON and markdown"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Save JSON (full data)
        json_path = f"data/matches_{timestamp}.json"
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2, ensure_ascii=False)
        print(f"   📄 JSON saved to {json_path}")
        
        # Save Markdown report (human-readable)
        md_path = f"data/matches_{timestamp}.md"
        self._generate_markdown_report(results, md_path)
        print(f"   📄 Report saved to {md_path}")
    
    def _generate_markdown_report(self, results: List[Dict], filepath: str):
        """Generate a nice markdown report of matches"""
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("# WaterlooWorks Job Matches Report\n\n")
            f.write(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write(f"**Total Jobs Analyzed:** {len(results)}\n\n")
            f.write("---\n\n")
            
            # Group by score ranges
            excellent = [r for r in results if r["match"]["fit_score"] >= 70]
            good = [r for r in results if 50 <= r["match"]["fit_score"] < 70]
            moderate = [r for r in results if 30 <= r["match"]["fit_score"] < 50]
            
            f.write(f"## Summary\n\n")
            f.write(f"- 🟢 **Excellent Matches (70+):** {len(excellent)}\n")
            f.write(f"- 🟡 **Good Matches (50-69):** {len(good)}\n")
            f.write(f"- 🟠 **Moderate Matches (30-49):** {len(moderate)}\n\n")
            f.write("---\n\n")
            
            # Write each job
            for i, result in enumerate(results, 1):
                job = result["job"]
                match = result["match"]
                
                # Emoji based on score
                if match["fit_score"] >= 70:
                    emoji = "🟢"
                elif match["fit_score"] >= 50:
                    emoji = "🟡"
                elif match["fit_score"] >= 30:
                    emoji = "🟠"
                else:
                    emoji = "🔴"
                
                f.write(f"## {emoji} #{i} - {job.get('title', 'Unknown')} ({match['fit_score']}/100)\n\n")
                f.write(f"**Company:** {job.get('company', 'N/A')}\n\n")
                f.write(f"**Location:** {job.get('location', 'N/A')}\n\n")
                
                # Match breakdown
                f.write(f"**Match Breakdown:**\n")
                f.write(f"- Keyword Match: {match.get('keyword_match', 0)}%\n")
                f.write(f"- Semantic Coverage: {match['coverage']}%\n")
                f.write(f"- Semantic Strength: {match['skill_match']}%\n")
                f.write(f"- Seniority Alignment: {match['seniority_alignment']}%\n\n")
                
                # Technology matching
                if match.get("matched_technologies"):
                    f.write(f"**✅ Matched Technologies ({len(match['matched_technologies'])}):**\n")
                    f.write(", ".join(match["matched_technologies"]) + "\n\n")
                
                if match.get("missing_technologies"):
                    f.write(f"**❌ Missing Technologies ({len(match['missing_technologies'])}):**\n")
                    f.write(", ".join(match["missing_technologies"]) + "\n\n")
                
                # Matched bullets
                if match.get("matched_bullets"):
                    f.write(f"**Your Relevant Experience ({len(match['matched_bullets'])} bullets):**\n\n")
                    for bullet in match["matched_bullets"][:5]:  # Top 5
                        f.write(f"- [{bullet['similarity']:.2f}] {bullet['text']}\n")
                    f.write("\n")
                
                # Job details
                if job.get("summary") or job.get("responsibilities"):
                    f.write(f"**Job Details:**\n\n")
                    if job.get("summary") and job["summary"] != "N/A":
                        f.write(f"{job['summary'][:300]}...\n\n")
                    if job.get("responsibilities") and job["responsibilities"] != "N/A":
                        f.write(f"**Responsibilities:** {job['responsibilities'][:200]}...\n\n")
                
                # Work arrangements
                if job.get("employment_location_arrangement") and job["employment_location_arrangement"] != "N/A":
                    f.write(f"**Location Arrangement:** {job['employment_location_arrangement']}\n\n")
                if job.get("work_term_duration") and job["work_term_duration"] != "N/A":
                    f.write(f"**Work Term:** {job['work_term_duration']}\n\n")
                
                if job.get("url"):
                    f.write(f"**Apply:** {job['url']}\n\n")
                
                f.write("---\n\n")
    
    def _show_summary(self, results: List[Dict]):
        """Show summary in terminal"""
        print("=" * 70)
        print("📊 TOP MATCHES")
        print("=" * 70)
        print()
        
        if not results:
            print("No matches found after filtering.")
            return
        
        # Show top 10
        for i, result in enumerate(results[:10], 1):
            job = result["job"]
            match = result["match"]
            
            # Emoji based on score
            if match["fit_score"] >= 70:
                emoji = "🟢"
            elif match["fit_score"] >= 50:
                emoji = "🟡"
            elif match["fit_score"] >= 30:
                emoji = "🟠"
            else:
                emoji = "🔴"
            
            print(f"{emoji} #{i} - Fit Score: {match['fit_score']}/100")
            print(f"   📋 {job.get('title', 'Unknown')} at {job.get('company', 'N/A')}")
            print(f"   📍 {job.get('location', 'N/A')}")
            print(f"   📈 Keyword: {match.get('keyword_match', 0)}% | Semantic: {match['coverage']}% | Seniority: {match['seniority_alignment']}%")
            
            # Show matched tech (if any)
            if match.get("matched_technologies"):
                tech_list = match["matched_technologies"][:5]  # Top 5
                print(f"   ✅ Tech Match: {', '.join(tech_list)}")
            
            print()
        
        if len(results) > 10:
            print(f"... and {len(results) - 10} more jobs")
            print()
        
        print("=" * 70)
        print(f"✅ Analysis complete! Check data/ folder for full reports.")
        print("=" * 70)


def main():
    """Main entry point"""
    import argparse
    
    parser = argparse.ArgumentParser(description="Analyze WaterlooWorks jobs")
    parser.add_argument(
        "--mode",
        choices=["realtime", "batch", "analyze", "cover-letter", "upload-covers", "apply", "scrape-folder", "generate-folder-covers"],
        default="batch",
        help="Operation mode: realtime (scrape+analyze live), batch (scrape then analyze), analyze (cached only), cover-letter (generate covers for Geese jobs), upload-covers (upload all covers to WW), apply (auto-apply with rules to any folder), scrape-folder (scrape and categorize jobs from a specific folder), generate-folder-covers (generate covers for folder-scraped jobs)"
    )
    parser.add_argument(
        "--quick",
        action="store_true",
        help="Quick mode (basic scraping, faster)"
    )
    parser.add_argument(
        "--force-rematch",
        action="store_true",
        help="Ignore cached matches and recalculate all job scores"
    )
    parser.add_argument(
        "--auto-save",
        action="store_true",
        help="Automatically save high-scoring jobs to WaterlooWorks folder"
    )
    parser.add_argument(
        "--max-apps",
        type=int,
        default=10,
        help="Max applications to submit in apply mode (default 10)"
    )
    parser.add_argument(
        "--skip-prescreening",
        action="store_true",
        help="Skip jobs with pre-screening questions in apply mode"
    )
    parser.add_argument(
        "--input-folder",
        type=str,
        default="input",
        help="Folder containing resume (default: input/)"
    )
    parser.add_argument(
        "--cover-letters-folder",
        type=str,
        default="cover_letters",
        help="Folder to save/read cover letters (default: cover_letters/)"
    )
    parser.add_argument(
        "--waterlooworks-folder",
        type=str,
        default=None,
        help="WaterlooWorks folder name to save/apply jobs to (default: from config.json). Use with --mode apply, scrape-folder, or generate-folder-covers"
    )
    
    args = parser.parse_args()
    
    analyzer = JobAnalyzer(
        input_folder=args.input_folder,
        cover_letters_folder=args.cover_letters_folder,
        waterlooworks_folder=args.waterlooworks_folder
    )
    
    if args.mode == "realtime":
        # Real-time processing mode
        analyzer.run_realtime_pipeline(auto_save_to_folder=True)
    
    elif args.mode == "analyze":
        # Load cached jobs and analyze
        print("📂 Using cached jobs...")
        with open("data/jobs_scraped.json", 'r', encoding='utf-8') as f:
            jobs = json.load(f)
        
        print(f"✅ Loaded {len(jobs)} jobs from cache\n")
        
        # Analyze
        print("🔍 Analyzing jobs...")
        if args.force_rematch:
            print("⚠️  Force rematch enabled - ignoring cache")
        results = analyzer.matcher.batch_analyze(jobs, force_rematch=args.force_rematch)
        
        # Filter and save
        filtered = analyzer._apply_filters(results)
        analyzer._save_results(filtered)
        analyzer._show_summary(filtered)
    
    elif args.mode == "cover-letter":
        # Generate cover letters for Geese Jobs
        print("=" * 70)
        print("📝 COVER LETTER GENERATOR")
        print("=" * 70)
        print()
        
        # Login with credentials from .env
        print("🔐 Logging into WaterlooWorks...")
        username = os.getenv("WATERLOOWORKS_USERNAME")
        password = os.getenv("WATERLOOWORKS_PASSWORD")
        
        if not username or not password:
            print("❌ Error: Credentials not found in .env file")
            print("   Please set WATERLOOWORKS_USERNAME and WATERLOOWORKS_PASSWORD")
            return
        
        print(f"   Using credentials from .env: {username}")
        print()
        
        auth = WaterlooWorksAuth(username, password)
        auth.login()
        
        # Get resume and prompt from config - use custom input folder
        default_resume = f"{analyzer.input_folder}/resume.pdf"
        resume_path = analyzer.config.get("resume_path", default_resume)
        
        # If config has "input/resume.pdf" but user specified custom folder, update path
        if resume_path == "input/resume.pdf" and analyzer.input_folder != "input":
            resume_path = default_resume
        
        cover_config = analyzer.config.get("cover_letter", {})
        template_path = cover_config.get("template_path", "template.docx")
        cached_jobs_path = cover_config.get("cached_jobs_path", "data/jobs_sample.json")
        prompt_template = cover_config.get("prompt_template", "Write a professional cover letter.")
        
        # Load resume text
        resume_text = ""
        
        # Try to load resume from different sources
        if os.path.exists(resume_path):
            # Load from PDF
            from pypdf import PdfReader
            reader = PdfReader(resume_path)
            resume_text = "".join(page.extract_text() for page in reader.pages)
            print(f"✓ Loaded resume from PDF: {resume_path}")
        elif os.path.exists("data/resume_parsed.txt"):
            # Fallback to parsed text
            with open("data/resume_parsed.txt", 'r', encoding='utf-8') as f:
                resume_text = f.read()
            print(f"✓ Loaded resume from text: data/resume_parsed.txt")
        else:
            print("❌ Error: Resume not found!")
            print(f"   Expected: {resume_path} OR data/resume_parsed.txt")
            print("   Please add your resume to one of these locations.")
            return
        
        # Initialize generator with custom cover letters folder
        from modules.cover_letter_generator import CoverLetterGenerator
        generator = CoverLetterGenerator(
            auth.driver, 
            resume_text, 
            prompt_template,
            api_key=os.getenv("GEMINI_API_KEY"),
            cover_letters_folder=analyzer.cover_letters_folder
        )
        
        # Generate all cover letters
        stats = generator.generate_all_cover_letters(cached_jobs_path)
        
        # Cleanup
        auth.driver.quit()
        
        # Show summary
        print("\n" + "=" * 70)
        print("📊 COVER LETTER GENERATION SUMMARY")
        print("=" * 70)
        print(f"Total Jobs: {stats['total_jobs']}")
        print(f"✅ Generated: {stats['generated']}")
        print(f"⏭  Skipped (already exists): {stats['skipped_existing']}")
        print(f"❌ Failed: {stats['failed']}")
        print("=" * 70)
    
    elif args.mode == "upload-covers":
        # Upload cover letters to WaterlooWorks
        print("=" * 70)
        print("📤 COVER LETTER UPLOADER")
        print("=" * 70)
        print()
        
        # Login with credentials from .env
        print("🔐 Logging into WaterlooWorks...")
        username = os.getenv("WATERLOOWORKS_USERNAME")
        password = os.getenv("WATERLOOWORKS_PASSWORD")
        
        if not username or not password:
            print("❌ Error: Credentials not found in .env file")
            print("   Please set WATERLOOWORKS_USERNAME and WATERLOOWORKS_PASSWORD")
            return
        
        print(f"   Using credentials from .env: {username}")
        print()
        
        auth = WaterlooWorksAuth(username, password)
        auth.login()
        
        # Initialize uploader with custom cover letters folder
        from modules.cover_letter_generator import CoverLetterUploader
        uploader = CoverLetterUploader(auth.driver, cover_letters_folder=analyzer.cover_letters_folder)
        
        # Upload all covers
        stats = uploader.upload_all_cover_letters()
        
        # Cleanup
        auth.driver.quit()
        
        # Show summary
        print("\n" + "=" * 70)
        print("📊 UPLOAD SUMMARY")
        print("=" * 70)
        print(f"Total Files: {stats['total_files']}")
        print(f"✅ Uploaded: {stats['uploaded']}")
        print(f"⏭️  Skipped (already uploaded): {stats['skipped_existing']}")
        print(f"❌ Failed: {stats['failed']}")
        print("=" * 70)
    
    elif args.mode == "apply":
        # Auto-apply to jobs in specified folder (default: Geese)
        print("=" * 70)
        print("🧑‍💻 WATERLOO WORKS APPLICATOR")
        print("=" * 70)
        print()

        # Determine folder name
        folder_name = args.waterlooworks_folder if args.waterlooworks_folder else analyzer.waterlooworks_folder
        print(f"📁 Target WaterlooWorks Folder: {folder_name}")
        print()

        # Load credentials from .env
        print("🔐 Logging into WaterlooWorks...")
        username = os.getenv("WATERLOOWORKS_USERNAME")
        password = os.getenv("WATERLOOWORKS_PASSWORD")
        if not username or not password:
            print("❌ Error: Credentials not found in .env file")
            print("   Please set WATERLOOWORKS_USERNAME and WATERLOOWORKS_PASSWORD")
            return
        print(f"   Using credentials from .env: {username}")
        print()

        # Load cached jobs to reference additional_info/external flags
        # Check for folder-specific jobs first, then fallback to default scraped jobs
        sanitized_folder = folder_name.replace(" ", "_").replace("/", "_").replace("\\", "_")
        folder_jobs_path = f"data/jobs_{sanitized_folder}.json"
        default_jobs_path = "data/jobs_scraped.json"
        
        if os.path.exists(folder_jobs_path):
            cached_path = folder_jobs_path
            print(f"📂 Loading jobs from folder-scraped file: {folder_jobs_path}")
        elif os.path.exists(default_jobs_path):
            cached_path = default_jobs_path
            print(f"📂 Loading jobs from default scraped file: {default_jobs_path}")
        else:
            print(f"❌ No jobs found! Tried:")
            print(f"   - {folder_jobs_path}")
            print(f"   - {default_jobs_path}")
            print(f"\n💡 Run scrape-folder mode first:")
            print(f'   python main.py --mode scrape-folder --waterlooworks-folder "{folder_name}"')
            return
        
        with open(cached_path, 'r', encoding='utf-8') as f:
            cached_jobs = json.load(f)
        
        print(f"✅ Loaded {len(cached_jobs)} jobs from {cached_path}")
        print()

        # Login and run applicator
        auth = WaterlooWorksAuth(username, password)
        auth.login()
        try:
            from modules.apply import WaterlooWorksApplicator
            applicator = WaterlooWorksApplicator(
                auth.driver,
                cover_letters_folder=analyzer.cover_letters_folder,
                waterlooworks_folder=folder_name
            )
            stats = applicator.apply_to_geese_jobs(
                cached_jobs, 
                max_applications=args.max_apps,
                skip_prescreening=args.skip_prescreening
            )
        finally:
            try:
                auth.close()
            except Exception:
                pass

        # Print summary
        print("\n" + "=" * 70)
        print("📊 APPLICATION SUMMARY")
        print("=" * 70)
        print(f"Attempted: {stats['attempted']} | Applied: {stats['applied']}")
        print(f"Skipped (extra docs): {len(stats['skipped_extra_docs'])}")
        print(f"Skipped (pre-screening): {len(stats['skipped_prescreening'])}")
        print(f"Skipped (missing cover letter): {len(stats['missing_cover_letter'])}")
        print(f"External apply needed: {len(stats['external_required'])}")
        print(f"Failed: {len(stats['failed'])}")

        if stats['missing_cover_letter']:
            print("\n📝 Skipped - Missing Cover Letter:")
            for jid, comp, title in stats['missing_cover_letter']:
                print(f" - [{jid}] {title} @ {comp}")
                print(f"   https://waterlooworks.uwaterloo.ca/myAccount/co-op/coop-postings.htm?ck_jobid={jid}")

        if stats['skipped_prescreening']:
            print("\n⏭️  Skipped due to pre-screening questions:")
            for jid, comp, title in stats['skipped_prescreening']:
                print(f" - [{jid}] {title} @ {comp}")
                print(f"   https://waterlooworks.uwaterloo.ca/myAccount/co-op/coop-postings.htm?ck_jobid={jid}")

        if stats['external_required']:
            print("\n🔗 Jobs requiring external application:")
            for jid, comp, title, hint in stats['external_required']:
                print(f" - [{jid}] {title} @ {comp}")
                print(f"   Hint: {hint}")
                print(f"   https://waterlooworks.uwaterloo.ca/myAccount/co-op/coop-postings.htm?ck_jobid={jid}")

        if stats['skipped_extra_docs']:
            print("\n📎 Skipped due to extra documents (beyond resume/cover):")
            for jid, comp, title, reason in stats['skipped_extra_docs']:
                print(f" - [{jid}] {title} @ {comp}")
                print(f"   Required: {reason}")
                print(f"   https://waterlooworks.uwaterloo.ca/myAccount/co-op/coop-postings.htm?ck_jobid={jid}")

        if stats['failed']:
            print("\n⚠️  Failures:")
            for jid, comp, title, reason in stats['failed']:
                print(f" - [{jid}] {title} @ {comp} -> {reason}")
        
        # Save detailed report to file
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_path = f"data/application_report_{timestamp}.md"
        
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write("# WaterlooWorks Application Report\n\n")
            f.write(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write("---\n\n")
            f.write("## Summary\n\n")
            f.write(f"- **Attempted:** {stats['attempted']}\n")
            f.write(f"- **Applied:** {stats['applied']}\n")
            f.write(f"- **Skipped (extra docs):** {len(stats['skipped_extra_docs'])}\n")
            f.write(f"- **Skipped (pre-screening):** {len(stats['skipped_prescreening'])}\n")
            f.write(f"- **Skipped (missing cover letter):** {len(stats['missing_cover_letter'])}\n")
            f.write(f"- **External apply needed:** {len(stats['external_required'])}\n")
            f.write(f"- **Failed:** {len(stats['failed'])}\n\n")
            f.write("---\n\n")
            
            if stats['missing_cover_letter']:
                f.write("## 📝 Skipped - Missing Cover Letter\n\n")
                f.write("*These jobs need cover letters to be generated/uploaded first.*\n\n")
                for jid, comp, title in stats['missing_cover_letter']:
                    f.write(f"### {title}\n")
                    f.write(f"**Company:** {comp}  \n")
                    f.write(f"**Job ID:** {jid}  \n")
                    f.write(f"**Action Needed:** Generate and upload cover letter, then apply manually  \n")
                    f.write(f"**Link:** [Apply Here](https://waterlooworks.uwaterloo.ca/myAccount/co-op/coop-postings.htm?ck_jobid={jid})\n\n")
            
            if stats['skipped_prescreening']:
                f.write("## ⏭️ Skipped - Pre-Screening Questions\n\n")
                f.write("*These jobs have pre-screening questions that need to be completed manually.*\n\n")
                for jid, comp, title in stats['skipped_prescreening']:
                    f.write(f"### {title}\n")
                    f.write(f"**Company:** {comp}  \n")
                    f.write(f"**Job ID:** {jid}  \n")
                    f.write(f"**Link:** [Apply Here](https://waterlooworks.uwaterloo.ca/myAccount/co-op/coop-postings.htm?ck_jobid={jid})\n\n")
            
            if stats['external_required']:
                f.write("## 🔗 External Application Required\n\n")
                f.write("*These jobs require applying through external portals or company websites.*\n\n")
                for jid, comp, title, hint in stats['external_required']:
                    f.write(f"### {title}\n")
                    f.write(f"**Company:** {comp}  \n")
                    f.write(f"**Job ID:** {jid}  \n")
                    f.write(f"**External Application Hint:** {hint}  \n")
                    f.write(f"**WaterlooWorks Link:** [View Job](https://waterlooworks.uwaterloo.ca/myAccount/co-op/coop-postings.htm?ck_jobid={jid})\n\n")
            
            if stats['skipped_extra_docs']:
                f.write("## 📎 Skipped - Additional Documents Required\n\n")
                f.write("*These jobs require extra documents beyond resume/cover letter (transcripts, portfolios, etc.).*\n\n")
                for jid, comp, title, reason in stats['skipped_extra_docs']:
                    f.write(f"### {title}\n")
                    f.write(f"**Company:** {comp}  \n")
                    f.write(f"**Job ID:** {jid}  \n")
                    f.write(f"**Required Document:** {reason}  \n")
                    f.write(f"**Link:** [Apply Here](https://waterlooworks.uwaterloo.ca/myAccount/co-op/coop-postings.htm?ck_jobid={jid})\n\n")
            
            if stats['failed']:
                f.write("## ⚠️ Failed Applications\n\n")
                f.write("*These applications encountered errors during the automation process.*\n\n")
                for jid, comp, title, reason in stats['failed']:
                    f.write(f"### {title}\n")
                    f.write(f"**Company:** {comp}  \n")
                    f.write(f"**Job ID:** {jid}  \n")
                    f.write(f"**Failure Reason:** {reason}  \n")
                    f.write(f"**Link:** [Try Again](https://waterlooworks.uwaterloo.ca/myAccount/co-op/coop-postings.htm?ck_jobid={jid})\n\n")
        
        print(f"\n📄 Detailed report saved to: {report_path}")
        print("=" * 70)
    
    elif args.mode == "scrape-folder":
        # Scrape jobs from a specific WaterlooWorks folder with intelligent categorization
        print("=" * 70)
        print("📂 FOLDER SCRAPER WITH REQUIREMENT ANALYSIS")
        print("=" * 70)
        print()
        
        folder_name = analyzer.waterlooworks_folder
        print(f"📁 Target folder: '{folder_name}'")
        
        # Check for existing scraped data to avoid re-scraping
        sanitized_folder = folder_name.replace(" ", "_").replace("/", "_")
        output_path = f"data/jobs_{sanitized_folder}.json"
        
        existing_jobs = {}
        if os.path.exists(output_path):
            try:
                with open(output_path, 'r', encoding='utf-8') as f:
                    existing_data = json.load(f)
                    # Convert to dict with job_id as key for fast lookup
                    existing_jobs = {job['id']: job for job in existing_data if job.get('id')}
                    print(f"📦 Loaded {len(existing_jobs)} existing jobs from cache")
            except Exception as e:
                print(f"⚠️  Could not load existing data: {e}")
        
        print()
        
        # Login with credentials from .env
        print("🔐 Logging into WaterlooWorks...")
        username = os.getenv("WATERLOOWORKS_USERNAME")
        password = os.getenv("WATERLOOWORKS_PASSWORD")
        
        if not username or not password:
            print("❌ Error: Credentials not found in .env file")
            print("   Please set WATERLOOWORKS_USERNAME and WATERLOOWORKS_PASSWORD")
            return
        
        print(f"   Using credentials from .env: {username}")
        print()
        
        auth = WaterlooWorksAuth(username, password)
        auth.login()
        
        try:
            from modules.folder_scraper import FolderScraper
            scraper = FolderScraper(
                auth.driver,
                api_key=os.getenv("GEMINI_API_KEY")
            )
            
            # Scrape and categorize jobs (with cache to skip already-scraped)
            jobs = scraper.scrape_folder_with_analysis(folder_name, existing_jobs=existing_jobs)
            
            # Save to structured file
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(jobs, f, indent=2, ensure_ascii=False)
            
            print("\n" + "=" * 70)
            print("📊 SCRAPING SUMMARY")
            print("=" * 70)
            print(f"Total Jobs: {len(jobs)}")
            
            # Count new vs cached
            new_count = sum(1 for j in jobs if j['id'] not in existing_jobs)
            cached_count = len(jobs) - new_count
            
            if cached_count > 0:
                print(f"🆕 New Jobs Scraped: {new_count}")
                print(f"📦 From Cache: {cached_count}")
            
            # Count by requirement type
            external = sum(1 for j in jobs if j.get('requirements', {}).get('external_application', {}).get('required'))
            extra_docs = sum(1 for j in jobs if j.get('requirements', {}).get('extra_documents', {}).get('required'))
            bonus = sum(1 for j in jobs if j.get('requirements', {}).get('bonus_items', {}).get('available'))
            standard = len(jobs) - external - extra_docs - bonus
            
            print(f"🔗 External Applications: {external}")
            print(f"📎 Extra Documents: {extra_docs}")
            print(f"⭐ Bonus Items Available: {bonus}")
            print(f"✅ Standard Jobs: {standard}")
            print()
            print(f"💾 Saved to: {output_path}")
            print("=" * 70)
            
        finally:
            try:
                auth.close()
            except Exception:
                pass

    elif args.mode == "generate-folder-covers":
        # Generate cover letters for folder-scraped jobs
        print("=" * 70)
        print("📝 FOLDER COVER LETTER GENERATOR")
        print("=" * 70)
        print()
        
        folder_name = analyzer.waterlooworks_folder
        sanitized_folder = folder_name.replace(" ", "_").replace("/", "_")
        jobs_file = f"data/jobs_{sanitized_folder}.json"
        
        # Check if jobs file exists
        if not os.path.exists(jobs_file):
            print(f"❌ Error: Jobs file not found: {jobs_file}")
            print(f"   Please run scrape-folder mode first:")
            print(f"   python main.py --mode scrape-folder --waterlooworks-folder \"{folder_name}\"")
            return
        
        # Load jobs
        with open(jobs_file, 'r', encoding='utf-8') as f:
            jobs = json.load(f)
        
        print(f"📁 Folder: '{folder_name}'")
        print(f"📦 Loaded {len(jobs)} jobs from {jobs_file}")
        print()
        
        # Get resume and config
        default_resume = f"{analyzer.input_folder}/resume.pdf"
        resume_path = analyzer.config.get("resume_path", default_resume)
        
        if resume_path == "input/resume.pdf" and analyzer.input_folder != "input":
            resume_path = default_resume
        
        cover_config = analyzer.config.get("cover_letter", {})
        template_path = cover_config.get("template_path", "template.docx")
        prompt_template = cover_config.get("prompt_template", "Write a professional cover letter.")
        
        # Load resume text
        resume_text = ""
        
        if os.path.exists(resume_path):
            from pypdf import PdfReader
            reader = PdfReader(resume_path)
            resume_text = "".join(page.extract_text() for page in reader.pages)
            print(f"✓ Loaded resume from PDF: {resume_path}")
        elif os.path.exists("data/resume_parsed.txt"):
            with open("data/resume_parsed.txt", 'r', encoding='utf-8') as f:
                resume_text = f.read()
            print(f"✓ Loaded resume from text: data/resume_parsed.txt")
        else:
            print("❌ Error: Resume not found!")
            print(f"   Expected: {resume_path} OR data/resume_parsed.txt")
            return
        
        # Configure Gemini
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            print("❌ Error: GEMINI_API_KEY not found in .env file")
            return
        
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel("gemini-2.0-flash-lite")
        
        # Create cover letters directory
        cover_letters_dir = Path(analyzer.cover_letters_folder)
        cover_letters_dir.mkdir(exist_ok=True)
        
        print()
        print("=" * 70)
        print("🔄 GENERATING COVER LETTERS")
        print("=" * 70)
        print()
        
        # Stats
        total = len(jobs)
        generated = 0
        skipped = 0
        failed = 0
        
        for i, job in enumerate(jobs, 1):
            company = job.get("company", "Unknown")
            title = job.get("title", "Unknown Position")
            job_id = job.get("id", "")
            
            print(f"[{i}/{total}] {title} @ {company}")
            
            # Generate sanitized filename (no illegal characters)
            def sanitize_filename(text):
                # Remove or replace Windows illegal characters
                text = text.replace('/', '_')
                text = text.replace('\\', '_')
                text = text.replace(':', '_')
                text = text.replace('*', '_')
                text = text.replace('?', '_')
                text = text.replace('"', '')
                text = text.replace('<', '')
                text = text.replace('>', '')
                text = text.replace('|', '_')
                text = text.replace('(', '')
                text = text.replace(')', '')
                text = text.replace('[', '')
                text = text.replace(']', '')
                text = text.replace('{', '')
                text = text.replace('}', '')
                # Remove other non-alphanumeric except spaces, hyphens, underscores
                text = re.sub(r'[^\w\s-]', '', text)
                # Replace multiple spaces with single space
                text = re.sub(r'\s+', ' ', text)
                # Replace spaces with underscores
                text = text.strip().replace(' ', '_')
                # Remove multiple underscores
                text = re.sub(r'_+', '_', text)
                return text.strip('_')
            
            # Sanitized filename (for saving)
            company_clean = sanitize_filename(company)
            title_clean = sanitize_filename(title)
            doc_name = f"{company_clean}_{title_clean}"
            pdf_path = cover_letters_dir / f"{doc_name}.pdf"
            
            # Store original name for metadata (will be used when uploading to WW)
            original_name = f"{company} - {title}"
            
            # Check if already exists
            if pdf_path.exists():
                print(f"         ⏭️  Already exists, skipping")
                skipped += 1
                continue
            
            # Retry logic for failed generations
            max_retries = 3
            retry_count = 0
            success = False
            last_error = None
            
            while retry_count < max_retries and not success:
                retry_count += 1
                try:
                    if retry_count > 1:
                        print(f"         🔄 Retry {retry_count-1}/{max_retries-1}...", end=" ")
                    
                    # Build description from job fields
                    description_parts = []
                    
                    if job.get("summary") and job["summary"] != "N/A":
                        description_parts.append(f"Job Summary:\n{job['summary']}")
                    
                    if job.get("responsibilities") and job["responsibilities"] != "N/A":
                        description_parts.append(f"\nResponsibilities:\n{job['responsibilities']}")
                    
                    if job.get("skills") and job["skills"] != "N/A":
                        description_parts.append(f"\nRequired Skills:\n{job['skills']}")
                    
                    if job.get("additional_info") and job["additional_info"] != "N/A":
                        description_parts.append(f"\nAdditional Info:\n{job['additional_info']}")
                    
                    description = "\n".join(description_parts) if description_parts else "No description available"
                    
                    # Generate cover letter with LLM
                    if retry_count == 1:
                        print(f"         🤖 Generating with AI...", end=" ")
                    
                    prompt = f"""{prompt_template}

{resume_text}

Job Information:
Organization: {company}
Position: {title}
Job Description: {description}

CRITICAL INSTRUCTIONS - READ CAREFULLY:
Write a professional cover letter (100-400 words) that is COMPLETELY FINISHED and READY TO SUBMIT.

ABSOLUTELY NO:
- Placeholders like [Your Name], [Company Name], [Position], etc.
- Square brackets [] with anything to fill in
- Blank spaces or underscores _____ to complete
- Notes like "customize this section" or "add details here"
- Any TODO items or suggestions for edits
- Special Unicode symbols, emojis, or decorative characters
- Smart quotes/apostrophes (use straight quotes: " and ')
- Em dashes or en dashes (use regular hyphens: -)
- Bullet points with special symbols (use plain text only)
- Any symbols that might not render properly in PDF

REQUIRED:
1. Use ONLY real information from the resume provided above
2. Reference the specific company name: {company}
3. Reference the specific position: {title}
4. Show genuine enthusiasm for THIS specific role at THIS specific company
5. Highlight 2-3 relevant experiences or skills from the resume that match the job description
6. Explain concretely why I'm a great fit based on my actual experience
7. Keep a professional but personable tone
8. Make every sentence complete and specific - no generic statements
9. Use only standard ASCII characters and basic punctuation (periods, commas, colons, semicolons, hyphens, parentheses)
10. Write in clear paragraphs with proper spacing

The cover letter will be exported as PDF and must be 100% ready to submit without any edits.

Start with "Dear Hiring Manager," and end with "Aman Zaveri"."""

                    response = model.generate_content(prompt)
                    cover_text = response.text.strip()
                    
                    # Clean up text
                    cover_text = re.sub(r"\[\[\d+\]\]", "", cover_text)  # Remove citation markers
                    
                    # Extract between markers
                    start = "Dear Hiring Manager,"
                    end = "Aman Zaveri"
                    idx1 = cover_text.find(start)
                    idx2 = cover_text.find(end)
                    
                    if idx1 != -1 and idx2 != -1:
                        cover_text = cover_text[idx1 + len(start) + 1 : idx2].strip()
                    
                    print("✓")
                    
                    # Save as PDF
                    print(f"         💾 Saving PDF...", end=" ")
                    
                    document = Document(template_path)
                    generated_text = document.add_paragraph().add_run(
                        f"""Dear Hiring Manager,
            
{cover_text}

Aman Zaveri"""
                    )
                    
                    generated_text.font.size = Pt(11)
                    generated_text.font.name = "Garamond"
                    
                    # Save as DOCX first
                    docx_path = cover_letters_dir / f"{doc_name}.docx"
                    document.save(str(docx_path))
                    
                    # Convert to PDF
                    pythoncom.CoInitialize()
                    convert(str(docx_path))
                    pythoncom.CoUninitialize()
                    
                    # Remove DOCX
                    docx_path.unlink()
                    
                    print(f"✓")
                    print(f"         ✅ {doc_name}.pdf")
                    generated += 1
                    success = True
                    
                except Exception as e:
                    last_error = e
                    if retry_count < max_retries:
                        print(f"❌ Error: {str(e)[:50]}")
                        import time
                        time.sleep(2)  # Wait before retry
                    else:
                        print(f"         ❌ Failed after {max_retries} attempts: {str(e)[:60]}")
                        failed += 1
            
            print()
        
        # Summary
        print("=" * 70)
        print("📊 GENERATION SUMMARY")
        print("=" * 70)
        print(f"Total Jobs: {total}")
        print(f"✅ Generated: {generated}")
        print(f"⏭️  Skipped (exists): {skipped}")
        print(f"❌ Failed: {failed}")
        print()
        print(f"📁 Cover letters saved to: {analyzer.cover_letters_folder}/")
        print("=" * 70)

    else:  # batch mode
        # Run full pipeline
        analyzer.run_full_pipeline(
            detailed=not args.quick, 
            force_rematch=args.force_rematch,
            auto_save_to_folder=args.auto_save
        )



if __name__ == "__main__":
    main()
